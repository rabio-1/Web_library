import os
from os import abort

from docx import Document
from flask import Flask, render_template, make_response, jsonify, request, flash
from data import db_session
from data.users import User
from data.books import Books
from forms.user import RegisterForm
from forms.book import BooksForm
from werkzeug.utils import redirect
from flask_login import LoginManager, login_user, logout_user, login_required, current_user

from forms.loginuser import LoginForm

app = Flask(__name__)
app.config['SECRET_KEY'] = 'yandexlyceum_secret_key'
login_manager = LoginManager()
login_manager.init_app(app)

# email1@mail.ru
# 12345678


@app.route("/")
def index():
    db_sess = db_session.create_session()
    books = db_sess.query(Books)
    return render_template("index.html", books=books)


@login_manager.user_loader
def load_user(user_id):
    db_sess = db_session.create_session()
    return db_sess.query(User).get(user_id)


@app.route('/login', methods=['GET', 'POST'])
def login():
    form = LoginForm()
    if form.validate_on_submit():
        db_sess = db_session.create_session()
        user = db_sess.query(User).filter(User.email == form.email.data).first()
        if user and user.check_password(form.password.data):
            login_user(user, remember=form.remember_me.data)
            return redirect("/")
        return render_template('login.html',
                               message="Неправильный логин или пароль",
                               form=form)
    return render_template('login.html', title='Авторизация', form=form)


@app.route('/register', methods=['GET', 'POST'])
def register():
    form = RegisterForm()
    if form.validate_on_submit():
        if form.password.data != form.password_again.data:
            return render_template('register.html', title='Регистрация',
                                   form=form,
                                   message="Пароли не совпадают")
        db_sess = db_session.create_session()
        if db_sess.query(User).filter(User.email == form.email.data).first():
            return render_template('register.html', title='Регистрация',
                                   form=form,
                                   message="Такой пользователь уже есть")
        user = User(
            name=form.name.data,
            email=form.email.data,
            about=form.about.data
        )
        user.set_password(form.password.data)
        db_sess.add(user)
        db_sess.commit()
        return redirect('/login')
    return render_template('register.html', title='Регистрация', form=form)


@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect("/")


@app.errorhandler(404)
def not_found(error):
    return make_response(jsonify({'error': 'Not found'}), 404)


@app.route('/books/new/', methods=['GET', 'POST'])
@login_required
def add_books():
    form = BooksForm()
    if form.validate_on_submit():
        db_sess = db_session.create_session()
        books = Books()
        books.title = form.title.data
        books.author = form.author.data
        books.content = form.content.data
        current_user.books.append(books)
        db_sess.merge(current_user)
        db_sess.commit()
        return redirect('/')
    return render_template('newBook.html', title='Добавление книги',
                           form=form)


@app.route('/books/<int:id>', methods=['GET', 'POST'])
@login_required
def edit_books(id):
    form = BooksForm()
    if request.method == "GET":
        db_sess = db_session.create_session()
        books = db_sess.query(Books).filter(Books.id == id,
                                            Books.user == current_user
                                            ).first()
        if books:
            form.title.data = books.title
            form.author.data = books.author
            form.content.data = books.content
        else:
            abort(404)
    if form.validate_on_submit():
        db_sess = db_session.create_session()
        books = db_sess.query(Books).filter(Books.id == id,
                                            Books.user == current_user
                                            ).first()
        if books:
            books.title = form.title.data
            books.author = form.author.data
            books.content = form.content.data
            db_sess.commit()
            return redirect('/')
        else:
            abort(404)
    return render_template('newBook.html',
                           title='Редактирование книги',
                           form=form
                           )


@app.route('/books_delete/<int:id>', methods=['GET', 'POST'])
@login_required
def books_delete(id):
    db_sess = db_session.create_session()
    books = db_sess.query(Books).filter(Books.id == id,
                                        Books.user == current_user
                                        ).first()
    if books:
        db_sess.delete(books)
        db_sess.commit()
    else:
        abort(404)
    return redirect('/')


@app.route('/books_read/<int:id>', methods=['GET'])
@login_required
def get_one_book(id):
    db_sess = db_session.create_session()
    books = db_sess.query(Books).get(id)
    return render_template("read.html", books=books)


@app.route('/books_download/<int:id>', methods=['GET', 'POST'])
@login_required
def download_book(id):
    db_sess = db_session.create_session()
    books = db_sess.query(Books).get(id)
    document = Document()
    document.add_heading(books.title, 0)
    document.add_heading(books.author, level=4)
    document.add_paragraph()
    document.add_paragraph(books.content)
    filename = f'{books.title}.docx'
    document.save('C:\\Users\\User\\Desktop\\' + filename)
    flash('Скачивание успешно')
    return redirect('/')


def main():
    db_session.global_init("db/books.db")
    app.run()


if __name__ == '__main__':
    main()